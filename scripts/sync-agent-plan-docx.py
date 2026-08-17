from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "AGENT-PROJECT-DOCUMENT-WORKFLOW-IMPLEMENTATION-PLAN.md"
OUTPUT = ROOT / "docs" / "AGENT-PROJECT-DOCUMENT-WORKFLOW-IMPLEMENTATION-PLAN.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TABLE_HEADER = "E8EEF5"
CODE_FILL = "F4F6F9"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = document.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_page(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "Agent 项目文档工作流实施计划"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8, color="6B7280")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer.add_run("第 ")
    set_run_font(label, size=8, color="6B7280")
    add_page_field(footer)
    suffix = footer.add_run(" 页")
    set_run_font(suffix, size=8, color="6B7280")


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^\)]+\))")


def add_inline(paragraph, text: str):
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position : match.start()]))
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
        elif token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True)
        else:
            label = token[1 : token.index("]")]
            target = token[token.index("(") + 1 : -1]
            run = paragraph.add_run(f"{label} ({target})")
            set_run_font(run, color=BLUE)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    weights = []
    for index in range(columns):
        length = max((len(row[index]) if index < len(row) else 0) for row in rows)
        weights.append(max(8, min(length, 60)))
    total = sum(weights)
    widths = [max(720, round(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    return widths


def configure_table(table, widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows


def add_table(document: Document, rows: list[list[str]]):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    widths = table_widths(rows)
    configure_table(table, widths)
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            value = row[column_index] if column_index < len(row) else ""
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, size=9, bold=row_index == 0)
            if row_index == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), TABLE_HEADER)
                cell._tc.get_or_add_tcPr().append(shd)


def convert():
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_styles(document)
    configure_page(document)

    in_code = False
    code_lines: list[str] = []
    first_heading = True
    index = 0
    while index < len(source_lines):
        line = source_lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph(style="Code Block")
                shade_paragraph(paragraph, CODE_FILL)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(source_lines) and source_lines[index].strip().startswith("|"):
                table_lines.append(source_lines[index].strip())
                index += 1
            add_table(document, parse_table(table_lines))
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if first_heading:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(12)
                run = paragraph.add_run(text)
                set_run_font(run, size=22, bold=True, color="0B2545")
                first_heading = False
            else:
                style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
                paragraph = document.add_paragraph(style=style)
                add_inline(paragraph, text)
            index += 1
            continue

        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            nesting = min(len(bullet.group(1)) // 2, 3)
            paragraph.paragraph_format.left_indent = Inches(0.375 + nesting * 0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            add_inline(paragraph, bullet.group(2))
            index += 1
            continue

        numbered = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            nesting = min(len(numbered.group(1)) // 2, 3)
            paragraph.paragraph_format.left_indent = Inches(0.375 + nesting * 0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            add_inline(paragraph, numbered.group(2))
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_after = Pt(8)
            shade_paragraph(paragraph, CODE_FILL)
            add_inline(paragraph, stripped.lstrip("> "))
            index += 1
            continue

        if stripped:
            paragraph = document.add_paragraph()
            if re.match(r"^(版本|日期|状态|适用范围)：", stripped):
                key, value = stripped.split("：", 1)
                set_run_font(paragraph.add_run(f"{key}："), bold=True, color=DARK_BLUE)
                add_inline(paragraph, value.rstrip())
            else:
                add_inline(paragraph, stripped.rstrip("  "))
        index += 1

    document.core_properties.title = "Agent 项目文档生成、审核与任务日志实施计划"
    document.core_properties.subject = "Provider 工具路由兼容性与 Agent 文档工作流实施计划"
    document.core_properties.comments = "由 Markdown 权威计划同步生成。"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    convert()
